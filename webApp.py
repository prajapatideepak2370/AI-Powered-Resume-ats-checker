import sys
# print(sys.version)
import os
import streamlit as st
from PIL import Image
from openai import OpenAI
from PyPDF2 import PdfReader
from docx import Document
import docx
import sys
# print(sys.version)

# ----------------------------- PAGE CONFIG -----------------------------------
img = Image.open(r"C:\\Users\\praja\\Downloads\\ChatGPT Image Aug 24, 2025, 11_22_50 PM.png")
st.set_page_config(page_title="CAREERSETU.AI", page_icon=img, layout="wide")

# ----------------------------- SESSION STATE INIT ----------------------------
if "started" not in st.session_state:
    st.session_state.started = False

# ----------------------------- OPENAI CLIENT --------------------------------
client = OpenAI(api_key=os.getenv("API_KEY"))

# ----------------------------- INTRO PAGE ------------------------------------
if not st.session_state.started:
    st.markdown(
        """
        <style>
        .intro-container {
            background: black;
            padding: 80px;
            border-radius: 16px;
            text-align: left;
        }
        .rainbow-text {
            font-size: 80px;
            font-weight: bold;
            background: linear-gradient(90deg, #ff0080, #ff8c00, #40e0d0, #8a2be2, #ff0080);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
        }
        .intro-sub {
            font-size: 46px;
            color: #111;
            font-weight: 600;
        }
        .intro-para {
            font-size: 25px;
            color: #333;
        }
        </style>
        <div class='intro-container'>
            <h1 class='rainbow-text'>CareerSetu.AI</h1>
            <h2 class='intro-sub'>AI Career Companion for Indian Graduates</h2>
            <p class='intro-para'>
                Tailor your resume, prepare for interviews, and map out your career path 
                with our bilingual AI platform.
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )

    if st.button("🚀 Get Started", key="start_btn"):
        st.session_state.started = True
        st.rerun()

# ----------------------------- HOME PAGE -------------------------------------
else:
    # -------- TOP NAVBAR --------
    nav1, nav2, nav3, nav4, nav5 = st.columns([6,.6,.5,.6,.5])
    with nav1:
        st.markdown("<h2 style='color:white;'>CAREERSETU.AI</h2>", unsafe_allow_html=True)
    with nav2:
        if st.button("Home"):
            st.info("Home page")
    with nav3:
        if st.button("More"):
            st.info("More Features")
    with nav4:
        if st.button("About"):
            st.info("About Section")
    with nav5:
        if st.button("Profile"):
            st.info("Profile Section")

    st.markdown("")
    st.markdown("")

    # -------- HERO SECTION --------
    st.markdown(
        """
        <h1 style='text-align: center; font-size: 62px; background: -webkit-linear-gradient(#60a5fa, #c084fc); 
        -webkit-background-clip: text; -webkit-text-fill-color: transparent;'>
            Smart Career Companion —
        </h1>
        <h2 style='text-align: center; font-size: 49px; background: -webkit-linear-gradient(#34d399, #22d3ee); 
        -webkit-background-clip: text; -webkit-text-fill-color: transparent;'>
            Powered by AI
        </h2>
        <p style='text-align: center; font-size: 22px; color: #d1d5db;'>
            Discover AI-driven tools to prepare, learn, and grow for placements & beyond.
        </p>
        """,
        unsafe_allow_html=True
    )

    # -------- FEATURE CARDS --------
    st.markdown("""
        <style>
        .card {
            background: linear-gradient(135deg, #1e293b, #0f172a);
            padding: 25px;
            border-radius: 12px;
            text-align: center;
            color: white;
            transition: 0.3s ease-in-out;
        }
        .card:hover {
            transform: translateY(-5px);
            background: linear-gradient(135deg, #3b82f6, #9333ea);
        }
        .icon {
            font-size: 40px;
            margin-bottom: 10px;
        }
        </style>
    """, unsafe_allow_html=True)

    # ----------------------------- PAGE ROUTING ----------------------------
    if "page" not in st.session_state:
        st.session_state.page = "home"


    # Helper to change page
    def go_to(page_name):
        st.session_state.page = page_name
        st.rerun()


    # ----------------------------- FEATURE CARDS ---------------------------
    if st.session_state.page == "home":
        st.markdown(""" 
            <style>
            .card {
                background: linear-gradient(135deg, #1e293b, #0f172a);
                padding: 25px;
                border-radius: 12px;
                text-align: center;
                color: white;
                transition: 0.3s ease-in-out;
            }
            .card:hover {
                transform: translateY(-5px);
                background: linear-gradient(135deg, #3b82f6, #9333ea);
            }
            .icon {
                font-size: 40px;
                margin-bottom: 10px;
            }
            </style>
        """, unsafe_allow_html=True)

        c1, c2, c3, c4 = st.columns(4)

        with c1:
            if st.button("📄 Resume Tailoring"):
                go_to("resume")
            
        with c2:
            if st.button("🎤 Interview Preparation"):
                go_to("interview")
        with c3:
            if st.button("📊 Placement Cell Mode"):
                go_to("placement")
        with c4:
            if st.button("🎓 Career Path Mapping"):
                go_to("career")

    # ----------------------------- INDIVIDUAL PAGES ------------------------
    
    # -----------------------------resume Tailoring ------------------------
    elif st.session_state.page == "resume":
        st.title("📄 Resume Tailoring")
        st.write("Upload your resume and get AI-powered improvement suggestions here.")

    # -------- Upload Resume --------
        uploaded = st.file_uploader("Upload Resume", type=["pdf","docx"])

    # -------- Job Description Input --------
        jd_text = st.text_area("Paste Job Description (Optional)", height=200)

    # -------- Resume Only Analysis --------
        if st.button("🔍 Analyze Resume"):
            if uploaded is None:
                st.warning("Please upload a resume file first.")
            else:
            # Extract text
                text = ""
                if uploaded.type == "application/pdf":
                    pdf = PdfReader(uploaded)
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                elif uploaded.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    doc = docx.Document(uploaded)
                    for para in doc.paragraphs:
                        text += para.text + "\n"

                if text.strip() == "":
                    st.error("Could not extract text from resume. Try another file.")
                else:
                    with st.spinner("Analyzing Resume with AI..."):
                        response = client.chat.completions.create(
                            model="gpt-4o-mini",
                            messages=[
                                {"role": "system", "content": "You are an expert career counselor."},
                                {"role": "user", "content": f"Analyze this resume and give strengths, weaknesses, and improvements:\n{text}"}
                            ]
                        )
                        result = response.choices[0].message.content
                    st.success(f"✅ Analysis Completed: {uploaded.name}")
                    st.markdown(result)

    # -------- JD Fit + ATS Score --------
        if st.button("📊 JD Fit & ATS Score"):
            if uploaded is None:
                st.warning("Please upload a resume file first.")
            elif jd_text.strip() == "":
                st.warning("Please paste a Job Description for JD Fit analysis.")
            else:
            # Extract text
                text = ""
                if uploaded.type == "application/pdf":
                    pdf = PdfReader(uploaded)
                    for page in pdf.pages:
                       text += page.extract_text() or ""
                elif uploaded.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    doc = docx.Document(uploaded)
                    for para in doc.paragraphs:
                        text += para.text + "\n"

                if text.strip() == "":
                    st.error("Could not extract text from resume. Try another file.")
                else:
                    with st.spinner("Comparing Resume with JD..."):
                        response = client.chat.completions.create(
                            model="gpt-4o-mini",
                            messages=[
                                {"role": "system", "content": "You are an ATS and career coach."},
                                {"role": "user", "content": f"""
                                Resume: {text}
                                Job Description: {jd_text}

                                1. Compare resume skills vs JD requirements.
                                2. Show Matching Skills and Missing Skills.
                                3. Give a JD Fit Score (0-100).
                                4. Simulate ATS Score with reasons.
                                """}
                           ]
                        )
                        result = response.choices[0].message.content
                    st.subheader("📊 JD Fit & ATS Analysis")
                    st.markdown(result)

        if st.button("⬅ Back to Home"):
            go_to("home")

#--------------------interview --------------------------------------------------------------------------

    elif st.session_state.page == "interview":
        st.title("🎤 Interview Preparation")
        st.write("Practice mock interviews with AI-driven feedback.")
        if st.button("⬅ Back to Home"):
            go_to("home")

#-------------------- Placement Cell Mode--------------------------------------------------------------------------

    elif st.session_state.page == "placement":
        st.title("📊 Placement Cell Mode")
        st.write("AI-powered tools to help you prepare for placements effectively.")

        option = st.radio("Select a Tool:", [
            "📌 Job Eligibility Checker",
            "📊 Company Insights",
            "📝 Skill Gap Analysis",
            "📈 Placement Statistics",
            "💡 Smart Recommendations"
        ])
        
    # -------- Job Eligibility Checker --------
        if option == "📌 Job Eligibility Checker":
            uploaded = st.file_uploader("Upload Resume", type=["pdf","docx"])
            jd_text = st.text_area("Paste Job Description", height=200)

            if st.button("🔍 Check Eligibility"):
                if not uploaded or not jd_text.strip():
                    st.warning("Upload resume & paste JD first.")
                else:
                    text = ""
                    if uploaded.type == "application/pdf":
                        pdf = PdfReader(uploaded)
                        for page in pdf.pages: text += page.extract_text() or ""
                    else:
                        doc = docx.Document(uploaded)
                        for para in doc.paragraphs: text += para.text + "\n"

                    with st.spinner("Checking eligibility..."):
                        response = client.chat.completions.create(
                            model="gpt-4o-mini",
                            messages=[
                                {"role": "system", "content": "You are an expert placement officer."},
                                {"role": "user", "content": f"Check if this resume fits the job description.\n\nResume:\n{text}\n\nJob Description:\n{jd_text}"}
                            ]
                        )
                        st.success("✅ Eligibility Report")
                        st.markdown(response.choices[0].message.content)
        # -------- Company Insights --------
        elif option == "📊 Company Insights":
            company = st.text_input("Enter Company Name")
            if st.button("🔍 Get Insights"):
                if company.strip() == "":
                   st.warning("Enter a company name.")
                else:
                    with st.spinner("Fetching insights..."):
                        response = client.chat.completions.create(
                            model="gpt-4o-mini",
                            messages=[
                                {"role": "system", "content": "You are a placement advisor with deep company insights."},
                                {"role": "user", "content": f"Provide placement-related insights about {company}. Include hiring trends, required skills, and interview tips."}
                            ]
                        )
                        st.subheader(f"📊 Insights for {company}")
                        st.markdown(response.choices[0].message.content)

    # -------- Skill Gap Analysis --------
        elif option == "📝 Skill Gap Analysis":
            uploaded = st.file_uploader("Upload Resume", type=["pdf","docx"], key="skill_gap_resume")
            jd_text = st.text_area("Paste Job Description", height=200, key="skill_gap_jd")

            if st.button("🔍 Analyze Skill Gaps"):
                if not uploaded or not jd_text.strip():
                    st.warning("Upload resume & paste JD first.")
                else:
                    text = ""
                    if uploaded.type == "application/pdf":
                        pdf = PdfReader(uploaded)
                        for page in pdf.pages: text += page.extract_text() or ""
                    else:
                        doc = docx.Document(uploaded)
                        for para in doc.paragraphs: text += para.text + "\n"

                    with st.spinner("Analyzing skill gaps..."):
                        response = client.chat.completions.create(
                            model="gpt-4o-mini",
                            messages=[
                                {"role": "system", "content": "You are a career coach specializing in skill analysis."},
                                {"role": "user", "content": f"Compare resume skills with this JD and suggest missing skills and learning roadmap.\n\nResume:\n{text}\n\nJob Description:\n{jd_text}"}
                            ]
                        )
                        st.subheader("📝 Skill Gap Report")
                        st.markdown(response.choices[0].message.content)

    # -------- Placement Statistics --------
        elif option == "📈 Placement Statistics":
            st.write("📊 Placement trends and mock statistics (AI-generated).")
            field = st.text_input("Enter Field / Role (e.g., Data Scientist, Software Engineer)")
            if st.button("📊 Show Stats"):
                if field.strip() == "":
                    st.warning("Enter a field or role.")
                else:
                    with st.spinner("Fetching placement stats..."):
                        response = client.chat.completions.create(
                            model="gpt-4o-mini",
                            messages=[
                                {"role": "system", "content": "You are an analyst providing placement statistics."},
                                {"role": "user", "content": f"Give placement statistics, average packages, and trends for {field} in India (mock data allowed)."}
                            ]
                        )
                        st.subheader(f"📈 Placement Stats for {field}")
                        st.markdown(response.choices[0].message.content)

    # -------- Smart Recommendations --------
        elif option == "💡 Smart Recommendations":
            uploaded = st.file_uploader("Upload Resume", type=["pdf","docx"], key="recommend_resume")
            if st.button("💡 Get Recommendations"):
                if not uploaded:
                    st.warning("Upload resume first.")
                else:
                    text = ""
                    if uploaded.type == "application/pdf":
                        pdf = PdfReader(uploaded)
                        for page in pdf.pages: text += page.extract_text() or ""
                    else:
                        doc = docx.Document(uploaded)
                        for para in doc.paragraphs: text += para.text + "\n"

                    with st.spinner("Generating recommendations..."):
                        response = client.chat.completions.create(
                            model="gpt-4o-mini",
                            messages=[
                                {"role": "system", "content": "You are a career mentor giving personalized recommendations."},
                                {"role": "user", "content": f"Based on this resume, suggest 3 suitable job roles, 3 online courses, and 2 internship ideas.\n\nResume:\n{text}"}
                            ] 
                        )
                        st.subheader("💡 Personalized Recommendations")
                        st.markdown(response.choices[0].message.content)

        if st.button("⬅ Back to Home"):
            go_to("home")       


    
    elif st.session_state.page == "career":
        st.title("🎓 Career Path Mapping")
        st.write("Personalized career roadmap & skill suggestions.")
        if st.button("⬅ Back to Home"):
            go_to("home")

    # -------- FOOTER --------
    st.markdown("###")
    st.caption("© 2025 CAREERSETU.AI")

    #------------Background ------
    st.markdown(
        """
        <style>
        .stApp {
            background: radial-gradient(circle at top left, #0f172a, #0a0a1f, #050510);
            color: white;
        }
        div.stButton > button:hover {
            background: linear-gradient(90deg, #f43f5e, #f97316, #facc15);
            transform: scale(1.08);
            box-shadow: 0 0 25px rgba(250, 204, 21, 0.8);
        }
        div.stButton > button:active {
            transform: scale(0.96);
            box-shadow: 0 0 18px rgba(37, 99, 235, 0.7);
        }
        </style>
        """,
        unsafe_allow_html=True
    )
